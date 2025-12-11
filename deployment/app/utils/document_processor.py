"""
Utilities for processing and extracting text from documents
"""
import os
from typing import Dict, Any, Optional, List, Union, Tuple
import logging
import asyncio
import uuid
import re
import json
from pathlib import Path
from datetime import datetime
from langchain_text_splitters.character import RecursiveCharacterTextSplitter
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
import typing
from typing import Iterator, List
from langchain_core.document_loaders import  BaseLoader
from langchain_core.documents import Document
from docling.document_converter import DocumentConverter
from docling.document_converter import DocumentConverter, PdfFormatOption, WordFormatOption
from docling.datamodel.base_models import InputFormat
from docling.pipeline.simple_pipeline import SimplePipeline
from docling.pipeline.standard_pdf_pipeline import StandardPdfPipeline
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
from langchain_text_splitters import MarkdownHeaderTextSplitter
headers_to_split_on = [
    ("#", "Header 1"),
    ("##", "Header 2"),
    ("###", "Header 3"),
]
def remove_image_tags(text: str) -> str:
    cleaned_text = re.sub(r'\s*<!--\s*image\s*-->\s*', ' ', text)
    return cleaned_text.strip()


def markdown_splitter(text: str, headers: List[Tuple[str, str]] = headers_to_split_on) -> List[str]:
    markdown_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers, strip_headers=False)
    docs = markdown_splitter.split_text(text)
    # goupping the docs by header with short sentences
    grouped_docs = []
    max_length = 1000
    current_length = 0
    content = ""
    while current_length < max_length:
        if not docs:
            break
        doc = docs.pop(0)
        doc_content = doc.page_content
        doc_metadata = doc.metadata
        if not doc_content:
            continue
        content += doc_content + "\n"
        current_length += len(doc_content)
        if current_length >= max_length or not docs:
            grouped_docs.append(Document(page_content=content, metadata=doc_metadata))
            content = ""
            current_length = 0
    return grouped_docs

def convert_document(file_path:str, type_doc:str = 'markdown'):

    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = False
    pipeline_options.do_table_structure = True
    document_coverter = (
        DocumentConverter(
            allowed_formats=[
                InputFormat.PDF,
                InputFormat.DOCX,
                InputFormat.HTML,
                InputFormat.PPTX,
                InputFormat.XLSX,
                InputFormat.CSV,
            ],
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options,
                                                             backend=PyPdfiumDocumentBackend),
                            InputFormat.DOCX: WordFormatOption(pipeline_cls = SimplePipeline)}
        )
    )

    coverted_docment  = document_coverter.convert(file_path , raises_on_error=False)
    if type_doc == 'markdown':
        return coverted_docment.document.export_to_markdown()
    else:
        return coverted_docment.document.export_to_dict()
# Setup logger
class DocumentPDFLoader(BaseLoader):

    def __init__(self, filepath: List[str]) -> None:
        self._filepath = filepath if isinstance(filepath, list) else [filepath]
        self._coverter = DocumentConverter()

    def lazy_load (self)->Iterator[Document]:
        for file in self._filepath:
            dl = self._coverter.convert(file).document
            text = dl.export_to_markdown()
            yield Document(page_content=text)

class DocumentCustomConverter(BaseLoader):

        def __init__(self, filepath: List[str], type_doc:str = 'markdown') -> None:
            self._filepath = filepath if isinstance(filepath, list) else [filepath]
            self._type_doc = type_doc
            self.pipeline_options = PdfPipelineOptions()
            self.pipeline_options.do_ocr = False
            self.pipeline_options.do_table_structure = True
            self.document_coverter = (
                DocumentConverter(
                    allowed_formats=[
                        InputFormat.PDF,
                        InputFormat.DOCX,
                        InputFormat.HTML,
                        InputFormat.PPTX,
                        InputFormat.XLSX,
                        InputFormat.CSV
                    ],
                    format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=self.pipeline_options,
                                                                    backend=PyPdfiumDocumentBackend),
                                    InputFormat.DOCX: WordFormatOption(pipeline_cls = SimplePipeline)}
                )
            )
        
        def lazy_load(self) -> Iterator[Document]:
            for file in self._filepath:
                converted_document  = self.document_coverter.convert(file , raises_on_error=False)
                if self._type_doc == 'markdown':
                    yield converted_document.document.export_to_markdown()
                else:
                    yield converted_document.document.export_to_text()

logger = logging.getLogger(__name__)

# Import environment variables
from app.core.config import settings

# Define vector store path
VECTOR_STORE_BASE_DIR = os.getenv("VECTOR_STORE_BASE_DIR", "./vector_stores_data")

# Document chunk size for vector storage
DEFAULT_CHUNK_SIZE = 1000
DEFAULT_CHUNK_OVERLAP = 200

def extract_text_from_document(file_path: str, file_type: str) -> Optional[str]:
    """
    Extract text from various document types for search indexing
    
    Args:
        file_path: Path to the document file
        file_type: Type of document (pdf, docx, txt, etc.)
        
    Returns:
        Extracted text or None if extraction failed
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return None
    
    try:
        if file_type.lower() == "pdf":
            return extract_text_from_pdf(file_path)
        elif file_type.lower() in ["doc", "docx"]:
            return extract_text_from_docx(file_path)
        elif file_type.lower() == "txt":
            return extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file type for text extraction: {file_type}")
            return None
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return None

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF files
    
    Requires PyPDF2 or pdfplumber to be installed:
    pip install PyPDF2
    or
    pip install pdfplumber
    """
    try:
        # Try PyPDF2 first
        import PyPDF2
        text = ""
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    except ImportError:
        try:
            # Fall back to pdfplumber
            import pdfplumber
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or "" + "\n"
            return text
        except ImportError:
            logger.error("Neither PyPDF2 nor pdfplumber are installed for PDF text extraction")
            return ""

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX files
    
    Requires python-docx to be installed:
    pip install python-docx
    """
    try:
        import docx
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except ImportError:
        logger.error("python-docx is not installed for DOCX text extraction")
        return ""

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text files"""
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    except UnicodeDecodeError:
        # Try another encoding if UTF-8 fails
        with open(file_path, "r", encoding="latin-1") as file:
            return file.read()

def update_document_metadata_with_text(doc_metadata: Dict[str, Any], extracted_text: str) -> Dict[str, Any]:
    """
    Update document metadata with extracted text for better search
    
    Args:
        doc_metadata: Existing document metadata dict
        extracted_text: Text extracted from document
        
    Returns:
        Updated metadata dict with extracted text
    """
    metadata = doc_metadata or {}
    
    # Add extracted text, but limit length to prevent metadata from getting too large
    max_text_length = 5000  # Limit to 5000 characters for performance reasons
    if extracted_text and len(extracted_text) > 0:
        metadata["extracted_text"] = extracted_text[:max_text_length]
        
        # Store word count
        word_count = len(extracted_text.split())
        metadata["word_count"] = word_count
        
        # Store page count if available in metadata
        if "page_count" not in metadata and "pages" in metadata:
            metadata["page_count"] = metadata["pages"]
    
    return metadata

def chunk_text(text: str, chunk_size: int = DEFAULT_CHUNK_SIZE, chunk_overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[str]:
    """
    Split text into overlapping chunks for processing by embedding models
    
    Args:
        text: The text to split into chunks
        chunk_size: Maximum chunk size in characters
        chunk_overlap: Overlap between chunks in characters
        
    Returns:
        List of text chunks
    """
    if not text:
        return []
        
    chunks = []
    for i in range(0, len(text), chunk_size - chunk_overlap):
        chunk = text[i:i + chunk_size]
        if chunk:
            chunks.append(chunk)
    
    return chunks

async def vectorize_document(
    document_id: int,
    text: str,
    metadata: Dict[str, Any],
    owner_type: str,
    owner_id: int,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP
) -> bool:
    """
    Process document text into vector store for RAG system
    
    Args:
        document_id: Database ID of the document
        text: Extracted text from document
        metadata: Document metadata
        owner_type: Type of document owner (customer, employee, guest)
        owner_id: ID of document owner
        chunk_size: Size of text chunks for embedding
        chunk_overlap: Overlap between chunks
        
    Returns:
        Success status of vectorization
    """
    try:
        # Create base directories if they don't exist
        vector_store_dir = os.path.join(VECTOR_STORE_BASE_DIR, owner_type, str(owner_id))
        os.makedirs(vector_store_dir, exist_ok=True)
        
        # Split text into chunks
        chunks = chunk_text(text, chunk_size, chunk_overlap)
        if not chunks:
            logger.warning(f"No text chunks extracted from document {document_id}")
            return False
            
        # Create vector store path for this document
        document_vector_path = os.path.join(vector_store_dir, f"doc_{document_id}")
        os.makedirs(document_vector_path, exist_ok=True)
        
        # Create chunks metadata
        chunks_data = []
        for i, chunk in enumerate(chunks):
            chunk_id = f"{document_id}_{i}"
            chunk_metadata = {
                "document_id": document_id,
                "chunk_id": chunk_id,
                "chunk_index": i,
                "total_chunks": len(chunks),
                **metadata
            }
            
            # Save chunk with metadata
            chunk_file = os.path.join(document_vector_path, f"chunk_{i}.json")
            with open(chunk_file, "w", encoding="utf-8") as f:
                json.dump({
                    "text": chunk,
                    "metadata": chunk_metadata
                }, f, ensure_ascii=False)
            
            chunks_data.append({
                "id": chunk_id,
                "path": chunk_file,
                "metadata": chunk_metadata
            })
        
        # Save document index file with references to all chunks
        index_file = os.path.join(document_vector_path, "index.json")
        with open(index_file, "w", encoding="utf-8") as f:
            json.dump({
                "document_id": document_id,
                "owner_type": owner_type,
                "owner_id": owner_id,
                "chunk_count": len(chunks),
                "chunks": chunks_data,
                "created_at": str(datetime.now())
            }, f, ensure_ascii=False)
            
        # Process vectors with embeddings if embeddings module is available
        try:
            from app.utils.embeddings import process_document_embeddings
            asyncio.create_task(process_document_embeddings(document_vector_path, chunks))
        except ImportError:
            logger.warning("Embeddings module not available, skipping vector processing")
        
        return True
    except Exception as e:
        logger.error(f"Error vectorizing document {document_id}: {str(e)}")
        return False

async def process_document_for_rag(
    document_id: int,
    file_path: str,
    file_type: str,
    metadata: Dict[str, Any],
    owner_type: str,
    owner_id: int
) -> bool:
    """
    Process a document for the RAG system - extract text and vectorize
    
    Args:
        document_id: Database ID of the document
        file_path: Path to the document file
        file_type: Type of document (pdf, docx, txt, etc.)
        metadata: Document metadata
        owner_type: Type of document owner (customer, employee, guest)
        owner_id: ID of document owner
        
    Returns:
        Success status of processing
    """
    # Extract text from document
    text = extract_text_from_document(file_path, file_type)
    if not text:
        logger.error(f"Failed to extract text from document {document_id}")
        return False
        
    # Update metadata with text preview
    updated_metadata = update_document_metadata_with_text(metadata, text)
    
    # Vectorize document
    success = await vectorize_document(
        document_id=document_id,
        text=text,
        metadata=updated_metadata,
        owner_type=owner_type,
        owner_id=owner_id
    )
    
    return success

def get_document_vector_path(document_id: int, owner_type: str, owner_id: int) -> str:
    """Get the path to the document's vector store directory"""
    return os.path.join(VECTOR_STORE_BASE_DIR, owner_type, str(owner_id), f"doc_{document_id}")

def check_document_in_vector_store(document_id: int, owner_type: str, owner_id: int) -> bool:
    """Check if document has been processed into the vector store"""
    vector_path = get_document_vector_path(document_id, owner_type, owner_id)
    index_file = os.path.join(vector_path, "index.json")
    return os.path.exists(index_file)

async def query_document_vectors(
    query: str,
    owner_type: Optional[str] = None,
    owner_id: Optional[int] = None,
    document_id: Optional[int] = None,
    limit: int = 5
) -> List[Dict[str, Any]]:
    """
    Query the vector store for relevant document chunks
    
    Args:
        query: The query text to search for
        owner_type: Optional filter by owner type
        owner_id: Optional filter by owner ID
        document_id: Optional filter by document ID
        limit: Maximum number of results to return
        
    Returns:
        List of relevant document chunks with metadata
    """
    try:
        # Import the embeddings module for vector search
        from app.utils.embeddings import search_embeddings
        
        # Determine search scope
        if document_id and owner_type and owner_id:
            # Search within a specific document
            search_path = get_document_vector_path(document_id, owner_type, owner_id)
            if not os.path.exists(search_path):
                return []
                
            return await search_embeddings(query, [search_path], limit)
            
        elif owner_type and owner_id:
            # Search within all documents for an owner
            search_path = os.path.join(VECTOR_STORE_BASE_DIR, owner_type, str(owner_id))
            if not os.path.exists(search_path):
                return []
                
            # Get all document directories
            doc_dirs = [os.path.join(search_path, d) for d in os.listdir(search_path) 
                       if os.path.isdir(os.path.join(search_path, d))]
            
            return await search_embeddings(query, doc_dirs, limit)
            
        else:
            # Search across all documents (limited to prevent performance issues)
            if not os.path.exists(VECTOR_STORE_BASE_DIR):
                return []
                
            # Get owner type directories
            owner_dirs = []
            for otype in os.listdir(VECTOR_STORE_BASE_DIR):
                otype_path = os.path.join(VECTOR_STORE_BASE_DIR, otype)
                if not os.path.isdir(otype_path):
                    continue
                    
                # Get owner directories
                for oid in os.listdir(otype_path):
                    owner_path = os.path.join(otype_path, oid)
                    if not os.path.isdir(owner_path):
                        continue
                        
                    # Add all document directories for this owner
                    doc_dirs = [os.path.join(owner_path, d) for d in os.listdir(owner_path)
                               if os.path.isdir(os.path.join(owner_path, d))]
                    owner_dirs.extend(doc_dirs)
            
            # Limit search to most recent 100 documents to prevent performance issues
            owner_dirs = sorted(owner_dirs, key=os.path.getctime, reverse=True)[:100]
            
            return await search_embeddings(query, owner_dirs, limit)
    
    except ImportError:
        logger.warning("Embeddings module not available for vector search")
        return []
    except Exception as e:
        logger.error(f"Error querying document vectors: {str(e)}")
        return []
